"""Genera Propuesta_Indigo_Decors.docx desde PROPUESTA_Odoo.md."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = Path(r"D:\01_Trabajo\Indigo\PROPUESTA_Odoo.md")
DST = Path(r"D:\01_Trabajo\Indigo\Propuesta_Indigo_Decors.docx")

INDIGO = RGBColor(0x1F, 0x44, 0x86)
NAVY = RGBColor(0x10, 0x1C, 0x36)
GRAY = RGBColor(0x55, 0x5B, 0x6E)
LIGHT_GRAY = "EEF1F7"

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# Base style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def parse_inline(paragraph, text):
    """Parse **bold** and `code` markers into runs."""
    # split keeping bold and inline code
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def add_heading(text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(22)
        run.font.color.rgb = INDIGO
    elif level == 2:
        run.font.size = Pt(16)
        run.font.color.rgb = INDIGO
    elif level == 3:
        run.font.size = Pt(13)
        run.font.color.rgb = NAVY
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = NAVY


def add_paragraph(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    parse_inline(p, text)


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    parse_inline(p, text)


def add_numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2)
    parse_inline(p, text)


def add_quote(lines):
    text = ' '.join(lines)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # left border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), '1F4486')
    pBdr.append(left)
    pPr.append(pBdr)
    parse_inline(p, text)
    for run in p.runs:
        run.italic = True
        run.font.color.rgb = GRAY


def add_code_block(lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run('\n'.join(lines))
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY


def add_table(rows):
    """rows: list of list[str], first row is header."""
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Light Grid Accent 1'
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            parse_inline(p, cell_text.strip())
            if r_idx == 0:
                set_cell_bg(cell, '1F4486')
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                if r_idx % 2 == 0:
                    set_cell_bg(cell, LIGHT_GRAY)


def parse_table(block_lines):
    """Parse a markdown table block into rows of cells."""
    rows = []
    for ln in block_lines:
        ln = ln.strip()
        if not ln or set(ln.replace('|', '').replace(':', '').replace('-', '').replace(' ', '')) == set():
            continue
        if re.match(r'^\|?[\s\-:|]+\|?$', ln):
            continue
        cells = [c.strip() for c in ln.strip('|').split('|')]
        rows.append(cells)
    return rows


# --- Parse markdown ---
lines = SRC.read_text(encoding='utf-8').splitlines()
i = 0
n = len(lines)

while i < n:
    line = lines[i]
    stripped = line.rstrip()

    # horizontal rule
    if stripped == '---':
        i += 1
        continue

    # heading
    m = re.match(r'^(#{1,6})\s+(.*)$', stripped)
    if m:
        level = len(m.group(1))
        add_heading(m.group(2).strip(), level)
        i += 1
        continue

    # code fence
    if stripped.startswith('```'):
        block = []
        i += 1
        while i < n and not lines[i].startswith('```'):
            block.append(lines[i])
            i += 1
        i += 1  # skip closing
        add_code_block(block)
        continue

    # blockquote (collect consecutive > lines)
    if stripped.startswith('>'):
        block = []
        while i < n and lines[i].lstrip().startswith('>'):
            block.append(lines[i].lstrip()[1:].strip())
            i += 1
        # collapse blank-line-separated paragraphs into one quote per block
        cur = []
        for ln in block:
            if ln == '':
                if cur:
                    add_quote(cur)
                    cur = []
            else:
                cur.append(ln)
        if cur:
            add_quote(cur)
        continue

    # table (line starts with | and next line looks like separator)
    if stripped.startswith('|') and i + 1 < n and re.match(r'^\|?[\s\-:|]+\|?$', lines[i + 1].strip()):
        block = []
        while i < n and lines[i].strip().startswith('|'):
            block.append(lines[i])
            i += 1
        rows = parse_table(block)
        if rows:
            add_table(rows)
        continue

    # bullet
    m = re.match(r'^[\-\*]\s+(.*)$', stripped)
    if m:
        # might be a list of multiple lines — but markdown bullets here are short
        add_bullet(m.group(1))
        i += 1
        # absorb continuation indented lines (rare in this doc)
        while i < n and re.match(r'^\s{2,}\S', lines[i]) and not re.match(r'^\s*[\-\*]\s', lines[i]):
            cont = lines[i].strip()
            doc.paragraphs[-1].add_run(' ' + cont)
            i += 1
        continue

    # numbered
    m = re.match(r'^\d+\.\s+(.*)$', stripped)
    if m:
        add_numbered(m.group(1))
        i += 1
        while i < n and re.match(r'^\s{3,}\S', lines[i]) and not re.match(r'^\s*\d+\.\s', lines[i]):
            cont = lines[i].strip()
            doc.paragraphs[-1].add_run(' ' + cont)
            i += 1
        continue

    # blank line
    if stripped == '':
        i += 1
        continue

    # regular paragraph — collect consecutive non-blank, non-special lines
    block = [stripped]
    i += 1
    while i < n:
        nxt = lines[i].rstrip()
        if nxt == '' or nxt == '---':
            break
        if re.match(r'^(#{1,6}\s|>|\||```|[\-\*]\s|\d+\.\s)', nxt):
            break
        block.append(nxt)
        i += 1
    add_paragraph(' '.join(block))

doc.save(DST)
print(f"OK -> {DST}")
